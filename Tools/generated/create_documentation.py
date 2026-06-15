# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

OUT = r"C:\Users\leona\Documents\PL Mathe\Abgabe\Dokumentation_Ballonfahrt_Leonardo_Freyhofer.docx"
G = r"C:\Users\leona\Documents\PL Mathe\Grafiken"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, size=8.5, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(1.25)
sec.bottom_margin = Cm(1.15)
sec.left_margin = Cm(1.35)
sec.right_margin = Cm(1.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(9.2)
normal.paragraph_format.space_after = Pt(3.2)
normal.paragraph_format.line_spacing = 1.04
for name, size, color in [("Heading 1", 13, "0B3D4A"), ("Heading 2", 10.5, "0B3D4A")]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(4)
    st.paragraph_format.space_after = Pt(2)


def para(text="", bold=False, size=None, color=None, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3.0)
    p.paragraph_format.line_spacing = 1.04
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.2)
    p.paragraph_format.space_after = Pt(1.8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("• " + text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(8.9)
    return p


para("Dokumentation zur Präsentationsleistung", bold=True, size=17, color="0B3D4A")
para("Ballonfahrt: Modellierung einer realistischen Horizontalgeschwindigkeit", bold=True, size=11.5, color="24515C")

meta = doc.add_table(rows=2, cols=4)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in meta.rows:
    for cell in row.cells:
        set_cell_border(cell)
        set_cell_shading(cell, "F7FAFC")
labels = ["Name", "Fach", "Thema", "Leitfrage"]
values = ["Leonardo Freyhofer", "Mathematik", "Funktion 3. Grades", "Wie entsteht ein realistisches Modell?"]
for i, txt in enumerate(labels):
    set_cell_text(meta.cell(0, i), txt, True, 8.0, "0B3D4A")
    set_cell_text(meta.cell(1, i), values[i], False, 8.0)

para("1. Aufgabenstellung und Modellidee", style="Heading 1")
para("Gesucht ist eine Funktion dritten Grades, die die Horizontalgeschwindigkeit einer siebenstündigen Ballonfahrt beschreibt. Die Fahrtrichtung soll sich zweimal umkehren, der Landepunkt soll horizontal möglichst nah am Startpunkt liegen und die Endgeschwindigkeit darf nicht 0 km/h sein.")
para("Für die Geschwindigkeit wird deshalb der Ansatz gewählt:")
para("v(t) = k · t · (t - r) · (t - s),      0 ≤ t ≤ 7", bold=True, size=10.5, color="0B3D4A", align=WD_ALIGN_PARAGRAPH.CENTER)
bullet("t = 0 beschreibt den Start der Fahrt.")
bullet("r und s sind die beiden Richtungswechsel, also Nullstellen innerhalb des Intervalls.")
bullet("k ist ein Streckfaktor und legt die Größenordnung der Geschwindigkeiten fest.")

para("2. Warum nicht die glatte Ausgangsvariante?", style="Heading 1")
para("Zunächst wirkt r = 3 attraktiv. Aus der Bedingung für den Endabstand ergibt sich dann s = 6,3. Wenn anschließend k so gewählt wird, dass die lokale Höchstgeschwindigkeit 25 km/h beträgt, entsteht jedoch v(7) ≈ 44,3 km/h. Diese Landegeschwindigkeit ist für eine Ballonfahrt zu hoch.")
img = os.path.join(G, "10_modellentscheidung_vergleich.png")
if os.path.exists(img):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img, width=Cm(16.4))
    p.paragraph_format.space_after = Pt(1)

para("3. Berechnung des zweiten Richtungswechsels", style="Heading 1")
para("Damit der Ballon horizontal wieder beim Startpunkt landet, muss die gesamte Ortsänderung null sein:")
para("∫₀⁷ v(t) dt = 0", bold=True, size=10.3, color="0B3D4A", align=WD_ALIGN_PARAGRAPH.CENTER)
para("Der Faktor k kann dabei zunächst weggelassen werden, weil er das Integral nur skaliert. Für t · (t-r) · (t-s) gilt:")
para("∫₀⁷ t(t-r)(t-s) dt = [t⁴/4 - (r+s)t³/3 + rs·t²/2]₀⁷ = 0", size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Wir wählen r = 3,4 h, damit die spätere Landegeschwindigkeit kleiner wird. Daraus folgt:")
para("s = ((343r/3) - (2401/4)) / ((49r/2) - (343/3)) ≈ 6,81579", bold=True, size=9.2, color="0B3D4A", align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

para("4. Finale Funktion und Kennwerte", style="Heading 1")
para("Der Streckfaktor wird nun so bestimmt, dass der lokale Hochpunkt genau 25 km/h beträgt. Damit ergibt sich als finales Modell:")
para("v(t) ≈ 1,648 · t · (t - 3,4) · (t - 6,81579)", bold=True, size=12, color="0B3D4A", align=WD_ALIGN_PARAGRAPH.CENTER)

vals = doc.add_table(rows=2, cols=5)
vals.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in vals.rows:
    for cell in row.cells:
        set_cell_border(cell)
for i, txt in enumerate(["Richtungswechsel", "Hochpunkt", "Tiefpunkt", "Landung", "Endabstand"]):
    set_cell_shading(vals.cell(0, i), "E8F2F5")
    set_cell_text(vals.cell(0, i), txt, True, 7.8, "0B3D4A")
for i, txt in enumerate(["3,4 h und 6,82 h", "25,0 km/h", "ca. -25,2 km/h", "v(7) ≈ 7,6 km/h", "0 km"]):
    set_cell_text(vals.cell(1, i), txt, False, 7.6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
speed_img = os.path.join(G, "01_geschwindigkeit_mit_flaechen.png")
if os.path.exists(speed_img):
    p.add_run().add_picture(speed_img, width=Cm(16.5))

para("5. Stammfunktion und räumliche Erweiterung", style="Heading 1")
para("Die Stammfunktion S(t)=∫v(t)dt beschreibt den horizontalen Abstand vom Startpunkt. Weil S(0)=0 und S(7)=0 gilt, liegt der Landepunkt horizontal wieder beim Startpunkt. Für die z-Achse wird eine Höhenfunktion ergänzt; die Zeit bleibt dabei nur der Parameter der Kurve.")

pic_table = doc.add_table(rows=1, cols=2)
pic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell in pic_table.rows[0].cells:
    set_cell_border(cell, "FFFFFF")
paths = [os.path.join(G, "02_stammfunktion_abstand.png"), os.path.join(G, "08_landeflaeche_gleichseitiges_dreieck_draufsicht.png")]
for cell, path in zip(pic_table.rows[0].cells, paths):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Cm(7.65))

para("Die Landefläche liegt in der Bodenebene z = 0. Sie wird als gleichseitiges Dreieck konstruiert; Start- und Landepunkt sind sein Zentrum und damit ausdrücklich kein Eckpunkt.", size=8.9)
para("Fazit: Das Modell erfüllt die Bedingungen der Aufgabe und ist realistischer als die glatte Ausgangsvariante, weil die Landegeschwindigkeit deutlich unter der sonstigen Höchstgeschwindigkeit liegt.", bold=True, size=9.1, color="0B3D4A")
para("Hilfsmittel: eigene Modellierung und Rechnung; Python/Matplotlib/Plotly zur Visualisierung; PowerPoint/Word zur Darstellung; ChatGPT als Arbeits- und Formulierungshilfe.", size=7.8, color="555555")

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.text = "Dokumentation Ballonfahrt-Modell – Leonardo Freyhofer"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(7.5)
    footer.runs[0].font.color.rgb = RGBColor.from_string("666666")

doc.save(OUT)
print(OUT)
